from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "manuscript" / "arxiv_manuscript.docx"
FIG = ROOT / "manuscript" / "figures"

TITLE = "Synthetic Exterior Point-Cloud Benchmarking for Rockpile Fragmentation: Fragment Generation, Rockpile Construction, Exterior-Only Learning, and EdgeConv P80 Estimation"
AUTHOR = "Namgyu (Noah) Ha"
AFFILIATION = "Department of Energy and Resources Engineering, National Korea Maritime and Ocean University, Busan, Republic of Korea"

ABSTRACT = (
    "Reliable rock-fragmentation monitoring is useful for blast assessment and mine-to-mill control, but real rockpile "
    "point clouds rarely provide point-level fragment labels or independent particle-size-distribution references. This "
    "paper presents a synthetic exterior point-cloud benchmark designed as a pre-field testbed for visible-surface PSD "
    "proxy estimation. The workflow starts from a library of synthetic rock-like fragment meshes with known volume and "
    "identity. These fragments are placed into labelled rockpile scenes, converted to exterior-only point-cloud scans, "
    "and used to train a graph edge-affinity model. The learning target is not a scene-specific fragment-ID class; "
    "instead, a DGCNN/EdgeConv model predicts whether neighbouring exterior points belong to the same fragment. Connected "
    "components of high-affinity edges become visible-surface clusters, and those clusters are converted into PSD and P80 "
    "proxy estimates. This formulation was chosen because fragment IDs are arbitrary across scenes, because local boundary "
    "decisions depend on neighbourhood geometry, and because hand-designed clustering rules such as DBSCAN or "
    "normal/curvature region growing struggle with irregular point density, touching particles, and over-merged surface "
    "patches. The latest benchmark run uses 100 no-boundary synthetic exterior scenes with 150 requested fragments per "
    "scene and a 60/20/20 train/validation/test split. The EdgeConv model reached validation average precision of 0.9313 "
    "after 24 epochs. However, validation selected a high edge threshold of 0.997, and the held-out test mean absolute P80 "
    "error was 19.04% with a mean noise fraction of 0.603. The result shows that EdgeConv learns useful same-fragment "
    "affinity, but it also makes clear that threshold calibration and connected-component post-processing remain the main "
    "obstacle before field deployment."
)

SECTIONS = [
    (
        "Introduction",
        [
            "Post-blast fragmentation affects loading, hauling, crushing, and comminution. Excessive fines or oversize material can lower equipment productivity and increase energy demand, so fragmentation monitoring is relevant to mine-to-mill decision support.",
            "Three-dimensional photogrammetry and lidar point clouds provide richer surface geometry than single images. They preserve relative position, local slope, and surface continuity. Nevertheless, a rockpile point cloud still observes only the exterior surface. Buried contacts and the back sides of fragments are unavailable. Consequently, a PSD estimated from a surface point cloud should be interpreted as a visible-surface proxy unless it is calibrated against independent field measurements.",
            "The benchmark reported here asks a narrow and reproducible question: from labelled synthetic rock fragments arranged into rockpiles, how much useful P80 proxy information can be recovered when the learning model sees only exterior points?",
        ],
    ),
    (
        "Related Work and Benchmark Positioning",
        [
            "Fragmentation measurement sits between blasting, materials handling, crushing, and sustainability. Blast outcomes influence downstream energy use, productivity, rehandling, and oversize management, so fragmentation monitoring is part of a broader mine-to-mill control problem rather than a purely visual segmentation problem.",
            "Image-based systems such as WipFrag and related digital image-processing approaches have been widely used because cameras are cheap and operationally convenient. Their limitation is not usefulness; it is the unavoidable ambiguity of inferring three-dimensional particle size from a two-dimensional visible surface.",
            "Three-dimensional reconstruction methods reduce some of that ambiguity. Structure-from-motion, multi-view stereo, and lidar workflows recover scale, surface orientation, and spatial context. Yet these data sources still observe only exposed surfaces. A surface scan cannot see buried material, so a fair benchmark must separate visible-surface proxy recovery from full-volume PSD measurement.",
            "Point-cloud learning offers useful tools for this exterior-surface problem, but the target must be chosen carefully. Classical partition metrics such as adjusted Rand index and normalized mutual information are useful diagnostics, but a P80 proxy can be sensitive to a small number of large over-merged or over-split components. The benchmark therefore reports both edge-learning behaviour and PSD proxy behaviour.",
        ],
    ),
    (
        "Materials and Methods",
        [
            "The methods are organised around the actual data path used in the benchmark: fragment meshes are generated first, fragments are arranged into labelled rockpile scenes, full mesh surfaces are sampled, hidden/contact samples are removed to create an exterior-only scan target, and only then are graph edges built for EdgeConv training.",
        ],
    ),
    (
        "Synthetic Fragment Mesh Library",
        [
            "The benchmark begins before pile construction. A library of rock-like triangular fragment meshes is generated with the Synthetic_Rockpile workflow. Each mesh has a persistent fragment identifier, volume, equivalent spherical diameter, and mesh path. This is important because the learning labels do not originate from a manual drawing on a point cloud; they originate from the synthetic fragment identity before the pile exists.",
            "For each scene, 150 fragments are sampled without replacement from the fragment catalogue. The reference PSD for a scene is computed from known fragment volumes using equivalent spherical diameter. This reference is independent of the model prediction, so the downstream P80 comparison tests the complete exterior-clustering pipeline rather than rewarding a method for reproducing its own clusters.",
        ],
    ),
    (
        "Rockpile Construction Backend",
        [
            "Several pile-generation options were tested during development, including the original Synthetic_Rockpile cone/drop heuristic, sequential-drop soft-sphere DEM, Project Chrono convex-hull contact, and clump contact. The fully dynamic sequential-drop variants were visually attractive but were too expensive or numerically sensitive for producing 100 repeatable no-boundary scenes.",
            "The reported dataset therefore uses the physics-informed realistic-rockpile generator in a no-boundary envelope-relax configuration with axis-clump contact. The method is best understood as a soft-sphere DEM surrogate for creating stable labelled piles, not as a calibrated FDEM blast simulation. Initial fragment centres are sampled inside a compact pile envelope. Gravity, damping, ground contact, and fragment-fragment contact are then integrated for a fixed number of steps. The visual geometry remains the original fragment mesh; the contact representation is an axis-clump approximation used for stable, faster collision response.",
        ],
    ),
    (
        "Material and Contact Interpretation",
        [
            "The density value, 2650 kg/m3, corresponds to a generic hard-rock density and is used to convert known mesh volumes into masses for the DEM surrogate. The contact parameters are numerical DEM parameters rather than laboratory-calibrated elastic moduli.",
            "This distinction matters for interpretation. A high-fidelity field simulator would require calibrated rock stiffness, fracture, shape distribution, bench geometry, blast throw, and post-blast interaction. The present generator is narrower: it provides labelled, repeatable exterior point clouds for training and testing segmentation-to-PSD algorithms.",
        ],
    ),
    (
        "Mesh Placement and Exterior Point-Cloud Sampling",
        [
            "After DEM relaxation, the original fragment meshes are placed at their final centres and orientations. Surface points are sampled from the placed meshes in proportion to fragment volume, with a minimum point allocation for small fragments. The full sampled pile still contains hidden surfaces: underside points, buried contact surfaces, and internal surfaces that would not be visible to a camera or lidar scanner.",
            "The model is not trained on the full mesh surface. It is trained on the exterior-only point-cloud target after visibility and height-envelope filtering. This distinction is central: the synthetic pile contains all labels, but the benchmark deliberately withholds non-visible geometry to avoid an unrealistically easy learning target.",
        ],
    ),
    (
        "Exterior-Only Scan Construction",
        [
            "The exterior scan construction is stricter than simple random surface sampling. Five default viewpoints are placed around the pile: four side viewpoints and one overhead viewpoint. For each viewpoint, points are binned in azimuth-elevation angle, and only the nearest point in each angular bin is retained. The angular bin width is 0.24 degrees in the reported dataset.",
            "The union of viewpoint-visible points is then filtered by a plan-view height envelope. Points are binned in the x-y plane using a 0.035 m grid. Within each grid cell, the local maximum elevation is estimated, and only points within 0.030 m of that upper envelope are retained. This second pass removes contact or interior samples that can otherwise survive through small gaps in point-sampled synthetic geometry.",
            "After exterior filtering, a small coordinate perturbation with standard deviation 0.0015 m is added, and a random density keep fraction between 0.78 and 0.92 is applied to mimic non-uniform point-cloud density. Normals and curvature are estimated from 30 nearest neighbours. A 12-nearest-neighbour graph is built on the exterior point coordinates, and each graph edge is labelled by same-fragment or different-fragment endpoint identity.",
        ],
    ),
    (
        "Dataset Summary",
        [
            "The latest production index contains 100 scenes with scene-level splitting: 60 training scenes, 20 validation scenes, and 20 held-out test scenes. The mean visible-fragment count after exterior filtering is 146.68, because a few requested fragments can be fully hidden or removed by the exterior filter.",
            "The mean exterior point count is 5735.63, the mean base radius is 0.898 m, and the mean pile height is 1.084 m. Scene-level splitting is used to avoid leakage: no points or graph edges from a training pile appear in validation or test scenes.",
        ],
    ),
    (
        "Why Edge Affinity and Why DGCNN/EdgeConv?",
        [
            "Direct point classification into fragment identifiers is poorly posed across scenes. Fragment 17 in one pile has no semantic relation to fragment 17 in another pile. A transferable model should therefore learn a relation, not a fixed label vocabulary. The benchmark uses edge affinity: for each edge in a local neighbourhood graph, the model predicts whether the two endpoints belong to the same underlying fragment.",
            "DGCNN/EdgeConv was selected because it learns local geometric relationships on point-cloud graphs. Compared with a global PointNet-style descriptor, EdgeConv better preserves local boundary information. Compared with PointNet++, the implementation is compact and directly suited to edge-pair supervision. Compared with an MLP edge classifier, EdgeConv embeds each point using graph neighbourhood context before scoring an edge. Compared with DBSCAN or normal/curvature region growing, it can learn how multiple cues trade off rather than relying on a fixed radius, angle, or smoothness threshold.",
        ],
    ),
    (
        "Training Protocol",
        [
            "Point features include normalized coordinates, normals, and curvature. Edge features encode local geometric differences and surface continuity cues. During training, edges are sampled in balanced positive/negative batches because same-fragment and different-fragment edges are naturally imbalanced in a local graph.",
            "The latest run trained the EdgeConv model for 24 epochs. The maximum number of sampled training edges per scene was 22,000, the maximum number of validation edges per scene was 36,000, and validation metrics were evaluated on 12 validation scenes per epoch for efficiency. Photogrammetry-realism augmentation was retained during training with strength 0.75.",
        ],
    ),
    (
        "PSD Proxy Evaluation",
        [
            "Predicted components are converted into visible-surface PSD proxies. For each predicted component, a surface-cluster diameter proxy is estimated from the spatial span of its points. Proxy volumes are then accumulated into a volume-weighted passing curve, from which P80 is interpolated.",
            "The evaluation reports both edge-learning and downstream PSD behaviour. Average precision and ROC-AUC measure how well the model ranks same-fragment edges. P80 error measures whether the resulting connected components preserve enough size information for a volume-weighted surface proxy. Noise fraction, NMI, and ARI diagnose whether predicted components resemble fragment instances.",
        ],
    ),
    (
        "Results",
        [
            "The EdgeConv model learned edge affinity clearly. Training loss decreased from 0.5722 at epoch 1 to 0.3479 at epoch 24. Validation average precision increased from 0.8569 to 0.9313, and validation ROC-AUC reached 0.9179.",
            "The downstream connected-component calibration remained difficult. The validation sweep selected the EdgeConv post-split variant at threshold 0.997. On the 20 held-out test scenes, this setting produced a mean absolute P80 error of 19.04%, a median absolute P80 error of 18.36%, and a mean noise fraction of 0.603.",
        ],
    ),
    (
        "Discussion",
        [
            "The result should be read carefully. The model is learning useful edge affinity, but the P80 pipeline is still limited by the conversion from probabilities to connected components. The selected threshold of 0.997 is a symptom of this calibration problem.",
            "This is why the benchmark is useful. Data generation is not the only bottleneck: the no-boundary 150-fragment scenes are stable and reproducible. Edge learning is not the only bottleneck either: validation AP is high. The remaining problem is graph partitioning and post-processing.",
        ],
    ),
    (
        "Limitations and Field Validation",
        [
            "This study reports a synthetic exterior-surface benchmark. It does not report field P80 accuracy. Real muckpile deployment requires independent references such as sieve or belt sampling, calibrated image-analysis outputs, carefully scaled manual annotation, or expert delineation on high-resolution orthomosaics.",
            "The synthetic piles remain simplified. The production preset was selected because it is stable and practical for 100 scenes, not because it fully reproduces blast mechanics. The present benchmark is therefore a pre-field testbed: it is useful for developing and comparing algorithms before expensive field validation.",
        ],
    ),
    (
        "Conclusion",
        [
            "Synthetic fragments are generated with known identity and volume, arranged into labelled rockpiles, converted to exterior-only point clouds, and used to train an EdgeConv edge-affinity model. DGCNN/EdgeConv is used because it learns local graph relationships that are more appropriate for arbitrary fragment IDs than direct point classification or purely geometric clustering. The latest 100-scene run confirms that edge affinity is learned well, with validation AP of 0.9313, but the held-out P80 result remains limited by high-threshold connected-component calibration.",
        ],
    ),
]

REFERENCES = [
    "Laurence, D. Establishing a sustainable mining operation: an overview. Journal of Cleaner Production, 2011.",
    "Cunningham, C. V. B. The Kuz-Ram model for prediction of fragmentation from blasting. 1st International Symposium on Rock Fragmentation by Blasting, 1983.",
    "Ouchterlony, F. The Swebrec function: linking fragmentation by blasting and crushing. Mining Technology, 2005.",
    "Palangio, T. C. et al. WIPFRAG: a breakthrough in fragmentation measurement. 21st Annual Conference on Explosives and Blasting Technique, 1995.",
    "Siddiqui, F. I. et al. Measurement of size distribution of blasted rock using digital image processing. Journal of King Abdulaziz University: Engineering Sciences, 2009.",
    "Westoby, M. J. et al. Structure-from-Motion photogrammetry: a low-cost, effective tool for geoscience applications. Geomorphology, 2012.",
    "Eltner, A. et al. Image-based surface reconstruction in geomorphometry: merits, limits and developments. Earth Surface Dynamics, 2016.",
    "Anderson, K. et al. Low-budget topography: Structure from Motion photogrammetry in geoscience applications. Progress in Physical Geography, 2019.",
    "Engin, I. C. et al. Practical measurement of size distribution of blasted rocks using LiDAR scan data. Rock Mechanics and Rock Engineering, 2020.",
    "Onederra, I. et al. Measuring blast fragmentation at Esperanza mine using high-resolution 3D laser scanning. Mining Technology, 2015.",
    "Wang, Y. et al. Dynamic Graph CNN for learning on point clouds. ACM Transactions on Graphics, 2019.",
    "Qi, C. R. et al. PointNet: Deep learning on point sets for 3D classification and segmentation. CVPR, 2017.",
    "Qi, C. R. et al. PointNet++: Deep hierarchical feature learning on point sets in a metric space. NeurIPS, 2017.",
    "Ester, M. et al. A density-based algorithm for discovering clusters in large spatial databases with noise. KDD, 1996.",
    "Hubert, L.; Arabie, P. Comparing partitions. Journal of Classification, 1985.",
    "Vinh, N. X. et al. Information theoretic measures for clusterings comparison. Journal of Machine Learning Research, 2010.",
    "Pedregosa, F. et al. Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 2011.",
    "Paszke, A. et al. PyTorch: An imperative style, high-performance deep learning library. NeurIPS, 2019.",
]


def set_styles(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color in [("Heading 1", 16, RGBColor(46, 116, 181)), ("Heading 2", 13, RGBColor(46, 116, 181))]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(6)


def add_centered(doc: Document, text: str, size: int, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold


def add_figure(doc: Document, filename: str, caption: str, width: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG / filename), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.paragraph_format.space_after = Pt(10)


def add_results_table(doc: Document) -> None:
    rows = [
        ("EdgeConv post split", "0.997", "19.04", "18.36", "0.603"),
        ("EdgeConv raw", "0.997", "20.82", "19.42", "0.603"),
        ("EdgeConv absorb + post split", "0.997", "21.84", "21.40", "0.506"),
        ("EdgeConv absorb", "0.997", "23.45", "21.40", "0.506"),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Variant", "Threshold", "Mean abs. P80 error (%)", "Median abs. P80 error (%)", "Noise fraction"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for par in cell.paragraphs:
            for run in par.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text
    cap = doc.add_paragraph("Table 3. Held-out test summary for the 20 test scenes.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True


def add_key_value_table(doc: Document, caption: str, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.0)
    for cell, text in zip(table.rows[0].cells, ["Item", "Value"]):
        cell.text = text
        for par in cell.paragraphs:
            for run in par.runs:
                run.bold = True
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.paragraph_format.space_after = Pt(10)


def add_pile_settings_table(doc: Document) -> None:
    add_key_value_table(
        doc,
        "Table 1. Production rockpile generation settings used for the 100-scene dataset.",
        [
            ("Placement backend", "physics-informed soft-sphere DEM surrogate"),
            ("Placement mode", "envelope_relax"),
            ("Contact model", "axis_clump"),
            ("Requested fragments per scene", "150"),
            ("Generated scenes", "100"),
            ("Scene split", "60 train / 20 validation / 20 test"),
            ("DEM steps and time step", "800 steps, dt = 5.0e-4 s"),
            ("Gravity", "9.81 m s-2"),
            ("Density", "2650 kg m-3"),
            ("Normal stiffness", "3.5e4 N m-1"),
            ("Normal / tangential damping", "115 / 65 N s m-1"),
            ("Fragment / ground friction coefficient", "0.88 / 1.05"),
            ("Collision radius scale", "0.55"),
            ("Clump radius scale / spread scale", "0.30 / 0.48"),
            ("Linear / angular damping", "0.028 / 0.018"),
            ("Temporary confinement wall", "none in production no-boundary run"),
        ],
    )


def add_exterior_settings_table(doc: Document) -> None:
    add_key_value_table(
        doc,
        "Table 2. Exterior scan and graph construction settings.",
        [
            ("Exterior viewpoint set", "four side viewpoints plus one overhead viewpoint"),
            ("Angular nearest-surface bin", "0.24 degrees"),
            ("Height-envelope grid", "0.035 m in plan view"),
            ("Height-envelope tolerance", "0.030 m below local maximum elevation"),
            ("Coordinate noise after filtering", "0.0015 m standard deviation"),
            ("Random density keep fraction", "uniform between 0.78 and 0.92"),
            ("Normal/curvature neighbourhood", "30 nearest neighbours"),
            ("Learning graph", "12-nearest-neighbour graph"),
            ("Edge label", "same-fragment vs different-fragment endpoint identity"),
        ],
    )


def build() -> None:
    doc = Document()
    set_styles(doc)
    add_centered(doc, TITLE, 16, bold=True)
    add_centered(doc, AUTHOR, 11)
    add_centered(doc, AFFILIATION, 10)
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(ABSTRACT)
    doc.add_paragraph("Keywords: rock fragmentation; synthetic rockpile; exterior point cloud; visible-surface PSD proxy; EdgeConv; DGCNN; P80")

    for title, paragraphs in SECTIONS:
        doc.add_heading(title, level=1)
        for text in paragraphs:
            doc.add_paragraph(text)
        if title == "Introduction":
            add_figure(doc, "workflow_schematic.png", "Figure 1. Benchmark workflow from synthetic fragments to exterior-only learning and field-validation roadmap.", 5.9)
        elif title == "Rockpile Construction Backend":
            add_figure(doc, "synthetic_generation_schematic.png", "Figure 2. Synthetic fragment, rockpile, and exterior-scan construction sequence.", 5.9)
            add_pile_settings_table(doc)
        elif title == "Mesh Placement and Exterior Point-Cloud Sampling":
            add_figure(doc, "rockpile_mesh_exterior_scene000.png", "Figure 3. Scene 000 mesh-level pile visualisation and exterior-only labelled point-cloud target.", 6.1)
        elif title == "Exterior-Only Scan Construction":
            add_exterior_settings_table(doc)
        elif title == "Dataset Summary":
            add_figure(doc, "dem_noboundary_relax150_scene000_preview.png", "Figure 4. Representative exterior point cloud from the no-boundary 150-fragment dataset.", 4.8)
        elif title == "Why Edge Affinity and Why DGCNN/EdgeConv?":
            add_figure(doc, "edge_affinity_schematic.png", "Figure 5. Edge-affinity formulation from local graph edges to connected components and PSD proxy.", 5.9)
        elif title == "Results":
            add_figure(doc, "02_edgeconv_training_curve.png", "Figure 6. Training loss and validation average precision for the 24-epoch EdgeConv run.", 5.4)
            add_results_table(doc)
            add_figure(doc, "03_edgeconv_test_p80_error_histogram.png", "Figure 7. Held-out test distribution of absolute P80 error.", 4.8)

    doc.add_heading("Acknowledgments", level=1)
    doc.add_paragraph("During preparation of this manuscript, the author used OpenAI Codex for coding, formatting, and drafting assistance. The author reviewed and edited the output and takes responsibility for the manuscript.")
    doc.add_heading("Data and Code Availability", level=1)
    doc.add_paragraph("Code, manuscript source, figures, and summary tables are prepared for release at https://github.com/Linkgyu/rockpile-edgeconv-arxiv. The full generated scenes can be regenerated using the included scripts.")
    doc.add_heading("References", level=1)
    for i, ref in enumerate(REFERENCES, start=1):
        doc.add_paragraph(f"{i}. {ref}")
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
