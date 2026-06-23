from __future__ import annotations

import csv
import math
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_DIR = ROOT / "manuscript"
FIG_DIR = MANUSCRIPT_DIR / "figures"
TABLE_DIR = ROOT / "results" / "tables"

OUT_MANUSCRIPT = MANUSCRIPT_DIR / "saimm_manuscript.docx"
OUT_COVER = MANUSCRIPT_DIR / "saimm_cover_letter.docx"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "0")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    for node in (fld_begin, instr, fld_sep, text, fld_end):
        r.append(node)


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.0))
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name, size, bold, italic, before, after in [
        ("Title", 15, True, False, 0, 8),
        ("Subtitle", 11, False, False, 0, 8),
        ("Heading 1", 12, True, False, 12, 6),
        ("Heading 2", 11, True, True, 8, 4),
        ("Heading 3", 11, False, True, 6, 3),
        ("Caption", 10, False, False, 3, 6),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if name == "Title":
            pPr = style._element.get_or_add_pPr()
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is not None:
                pPr.remove(pBdr)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def p(doc: Document, text: str = "", style: str | None = None):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if text:
        para.add_run(text)
    return para


def heading(doc: Document, text: str, level=1):
    return p(doc, text, f"Heading {level}")


def add_caption(doc: Document, label: str, caption: str):
    para = p(doc, style="Caption")
    run = para.add_run(f"{label}. ")
    run.bold = True
    para.add_run(caption.rstrip("."))
    return para


def add_image(doc: Document, filename: str, width_cm: float, caption: str, label: str):
    path = FIG_DIR / filename
    if path.exists():
        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.add_run().add_picture(str(path), width=Cm(width_cm))
        add_caption(doc, label, caption)
    else:
        add_caption(doc, label, f"[Figure file missing: {filename}] {caption}")


def table_from_rows(doc: Document, caption: str, rows, widths, label: str):
    add_caption(doc, label, caption)
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, val in enumerate(rows[0]):
        cell = hdr.cells[i]
        cell.text = str(val)
        set_cell_shading(cell, "EDEDED")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    set_table_widths(table, widths)
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing = 1.15
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                    run.font.size = Pt(9)
    p(doc)
    return table


def read_csv(path: Path):
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def fmt(x, digits=2):
    if x in ("", None):
        return "-"
    try:
        val = float(x)
    except (TypeError, ValueError):
        return str(x)
    if math.isnan(val):
        return "-"
    return f"{val:.{digits}f}"


def build_manuscript():
    doc = Document()
    configure_doc(doc)

    title = (
        "Synthetic exterior point-cloud benchmark for rockpile fragmentation: "
        "fragment generation, rockpile construction, EdgeConv learning, and size-distribution proxy evaluation"
    )
    tp = p(doc, title, "Title")
    tp.runs[0].bold = True

    p(doc, "N. Ha1", "Subtitle")
    p(doc, "1 Department of Energy and Resources Engineering, National Korea Maritime and Ocean University, Busan, Republic of Korea; ORCID: 0009-0007-4918-6239")
    p(doc, "Corresponding author: N. Ha, noah.ha.au@gmail.com")

    heading(doc, "Abstract")
    abstract = (
        "Post-blast rockpile fragmentation controls loading, hauling, and crushing performance, but field point clouds "
        "rarely provide fragment-level labels or independent size references. This study presents a synthetic benchmark "
        "for exterior rockpile point clouds in which individual rock-like fragments are generated first, assembled into "
        "labelled piles, filtered to retain only visible exterior points, and then used to compare fragment clustering "
        "methods. A graph EdgeConv model predicts whether neighbouring exterior points belong to the same fragment, and "
        "thresholded edge affinities are converted into connected components and visible-surface particle-size proxies. "
        "The same 100 scenes, each containing 150 fragments, were used for all methods with a 60:20:20 train, validation, "
        "and test split. The revised dataset applies multi-view Hidden Point Removal to the stored fragment poses, producing "
        "a stricter photogrammetry-like exterior shell. EdgeConv learned strong same-fragment affinity on this target, with "
        "validation average precision increasing to 0.962 after 24 epochs. The experiment is reported as three linked "
        "tasks: segmentation against exterior labels, exterior-size proxy recovery against oracle exterior labels, and "
        "full-pile PSD recovery after a validation-fitted exterior-to-full calibration layer. Oracle exterior labels still "
        "overestimated full-pile P80 by 22.94 per cent on the held-out test scenes, showing that exterior visibility imposes "
        "a systematic size bias even before model error is considered. After fitting a linear P80 calibration only on the "
        "validation scenes, EdgeConv reduced held-out full-P80 mean absolute error from 29.26-34.15 per cent to 5.00-5.46 "
        "per cent across the tested post-processing variants. The benchmark is intended as a reproducible pre-field testbed "
        "for separating visible-fragment segmentation from mine-site PSD calibration."
    )
    p(doc, abstract)
    p(doc, "Keywords: fragmentation monitoring; rockpile; point cloud segmentation; synthetic benchmark; EdgeConv; mine-to-mill")

    heading(doc, "Introduction")
    p(doc, (
        "Blast fragmentation affects digging, loading, hauling, crushing, and comminution. A reliable estimate of the "
        "coarse end of the fragment-size distribution can help engineers detect oversize material, adjust blast designs, "
        "and stabilise crusher feed. Two-dimensional image-analysis systems remain attractive because they are practical, "
        "but they must infer scale, overlap, and hidden fragment boundaries from a projected view. Three-dimensional "
        "photogrammetry and laser scanning preserve richer surface geometry, yet a scan still observes only the outside "
        "of the muckpile. Buried contacts and hidden backs of fragments remain unknown."
    ))
    p(doc, (
        "This study therefore treats surface estimates as visible-surface size-distribution proxies. The word proxy is "
        "important. The benchmark does not claim that an exterior point cloud directly measures the whole-pile particle "
        "size distribution. Instead, it asks how much reference size information can be recovered when the algorithm is "
        "given only the kind of exterior points that a camera or scanner could plausibly see."
    ))
    p(doc, (
        "The contribution is a reproducible synthetic pre-field workflow. Rock-like fragments are generated with known "
        "identities and volumes; piles are constructed from those fragments; full synthetic geometry is converted into "
        "exterior-only point clouds; and an edge-affinity model is compared with graph, density, region-growing, and "
        "shallow learning baselines. The manuscript emphasises the construction of the labelled piles and exterior "
        "learning target because these steps determine whether a method is being tested on a realistic scan proxy or on "
        "unobservable interior geometry."
    ))
    add_image(doc, "workflow_schematic.png", 15.8, "Benchmark workflow from synthetic fragment generation to exterior-only learning and size-distribution proxy evaluation", "Figure 1")

    heading(doc, "Related work and benchmark positioning")
    p(doc, (
        "Fragmentation models and monitoring systems have a long history in blasting and mine-to-mill optimisation "
        "(Cunningham, 1983; Ouchterlony, 2005). Commercial and research image-analysis systems such as WipFrag made "
        "fragmentation measurement operationally accessible, but their two-dimensional measurements remain sensitive to "
        "occlusion, scale, fines, and manual editing (Palangio et al., 1995; Siddiqui et al., 2009). Structure-from-motion "
        "and laser scanning provide point clouds that are more suitable for surface-shape reasoning, and have been applied "
        "in geoscience and mining contexts (Westoby et al., 2012; Eltner et al., 2016; Engin et al., 2020)."
    ))
    p(doc, (
        "Point-cloud neural networks including PointNet, PointNet++, and dynamic graph convolutional neural networks "
        "show that local geometric relationships can be learned directly from irregular samples (Qi et al., 2017a; "
        "Qi et al., 2017b; Wang et al., 2019). For rockpiles, however, a direct fragment-ID classifier is poorly posed "
        "because fragment labels are arbitrary from one scene to another. Edge affinity is a more transferable target: "
        "neighbouring points are classified as same-fragment or different-fragment, after which connected components "
        "form candidate visible fragments."
    ))

    heading(doc, "Materials and methods")
    heading(doc, "Synthetic fragment mesh library", 2)
    p(doc, (
        "The benchmark begins with a library of synthetic rock-like fragment meshes. Each mesh carries a fixed fragment "
        "identity, volume, equivalent spherical diameter, and local surface samples. Fragment generation is kept separate "
        "from rockpile construction so that the size reference is known before any pile placement, exterior filtering, or "
        "machine-learning step is applied. This prevents the downstream segmentation algorithm from defining its own "
        "reference size distribution."
    ))
    add_image(doc, "synthetic_generation_schematic.png", 15.8, "Synthetic fragment and rockpile generation sequence with label-preserving mesh placement and exterior sampling", "Figure 2")

    heading(doc, "Rockpile construction backend", 2)
    p(doc, (
        "The production dataset used the no-boundary envelope-relax and axis-clump backend developed after earlier "
        "sequential-drop and Chrono trials. Those trials were useful for diagnosing contact explosions, excessive lateral "
        "dispersion, and boundary-box artefacts, but the final reported 100-scene dataset was produced by a deterministic "
        "CPU-parallel placement workflow that preserved individual fragment meshes and stored the scene centres and "
        "orientations required for reconstruction."
    ))
    table_from_rows(
        doc,
        "Production rockpile generation settings used for the 100-scene dataset",
        [
            ["Item", "Value"],
            ["Number of scenes", "100"],
            ["Fragments per scene", "150"],
            ["Scene split", "60 training, 20 validation, 20 test scenes"],
            ["Pile backend", "No-boundary envelope-relax placement with axis-clump contact approximation"],
            ["Boundary condition", "Open floor; no confinement box in the final production run"],
            ["Placement objective", "Compact mound-like pile with overlap rejection and radial containment penalty"],
            ["Stored outputs", "Per-fragment centre, orientation, mesh identity, labels, and exterior point-cloud samples"],
        ],
        [3000, 6360],
        "Table 1",
    )

    heading(doc, "Material and contact interpretation", 2)
    p(doc, (
        "The material parameters were selected to represent dense hard rock rather than lightweight artificial particles. "
        "Because the production backend is a geometry-preserving placement and relaxation model rather than a full field "
        "calibration of breakage mechanics, the parameters should be read as physically informed contact settings used to "
        "stabilise pile morphology and prevent unrealistically elastic scattering."
    ))
    table_from_rows(
        doc,
        "Material and contact parameters used to guide the hard-rock pile-generation trials",
        [
            ["Parameter", "Value"],
            ["Density", "2670 kg/m3"],
            ["Young's modulus", "5.0e10 Pa"],
            ["Poisson ratio", "0.25"],
            ["Coefficient of restitution", "0.20"],
            ["Static friction coefficient", "0.85"],
            ["Sliding friction coefficient", "0.65"],
            ["Rolling friction coefficient", "0.15"],
            ["Spinning friction coefficient", "0.03"],
        ],
        [3600, 5760],
        "Table 2",
    )

    heading(doc, "Mesh placement and exterior point-cloud sampling", 2)
    p(doc, (
        "After placement, source fragment meshes were reconstructed at their stored scene centres and orientations. This "
        "mesh-level reconstruction was used for visual audit and for confirming that fragment identity had not been lost "
        "during scene construction. The learning data were not sampled from arbitrary internal surfaces. Instead, the full "
        "mesh scene was converted into an exterior scan proxy using multi-view Hidden Point Removal. This change preserves "
        "the stored rockpile placement while replacing the earlier plan-view height-envelope filter, which could remove "
        "legitimate side-visible exterior points."
    ))
    add_image(doc, "rockpile_mesh_exterior_scene000.png", 15.8, "Scene 000 mesh-level pile visualisation and exterior-only point-cloud target with retained fragment labels", "Figure 3")
    add_image(doc, "dem_noboundary_relax150_scene000_preview.png", 14.8, "Representative exterior point cloud from the no-boundary 150-fragment dataset", "Figure 4")

    heading(doc, "Exterior-only scan construction", 2)
    table_from_rows(
        doc,
        "Exterior scan and graph construction settings used for the learning target",
        [
            ["Stage", "Setting"],
            ["Surface visibility", "Hidden Point Removal using spherical flipping and convex-hull visibility"],
            ["Viewpoints", "Eight side-ring viewpoints plus one overhead viewpoint"],
            ["Envelope filtering", "None in the production HPR run; side-visible exterior points are retained"],
            ["Label transfer", "Each retained point stores the source fragment identity"],
            ["Graph construction", "Local point graph built from exterior samples only"],
            ["Learning target", "Binary same-fragment edge affinity"],
            ["Post-processing", "Thresholded affinities converted to connected components and size proxies"],
        ],
        [3000, 6360],
        "Table 3",
    )
    p(doc, (
        "The exterior filter is deliberately conservative. A full synthetic pile includes contact patches and interior "
        "faces that cannot be observed from a field scan. Training on those points would make the benchmark easier but "
        "less relevant. The HPR exterior-only construction therefore sacrifices some label completeness to preserve the "
        "visibility constraint. Across the 100 HPR-refiltered scenes, the mean retained fragment count was 141.89 out of "
        "150 requested fragments, the mean exterior point count was 6975.05, and the mean local-graph positive edge "
        "fraction was 0.924."
    ))
    add_image(doc, "exterior_filter_section_scan_scene000.png", 15.8, "Exterior-only HPR diagnostic for scene 000. The half-section shows that hidden interior samples are removed while the retained scan follows the visible outside shell", "Figure 5")

    heading(doc, "Edge-affinity learning and baseline methods")
    heading(doc, "Why EdgeConv was used", 2)
    p(doc, (
        "EdgeConv was selected because the task is local and relational. Rockpile segmentation depends on whether adjacent "
        "surface patches belong to the same fragment, and this decision is influenced by relative position, local surface "
        "shape, and neighbourhood context. A scene-level fragment-ID classifier would not transfer across piles because "
        "fragment numbers are arbitrary. A purely global network would also discard the local contact information that "
        "separates neighbouring stones. EdgeConv instead forms local graph features and updates point representations "
        "from neighbouring relationships, matching the edge-affinity formulation of the benchmark."
    ))
    add_image(doc, "edge_affinity_schematic.png", 15.0, "Edge-affinity formulation from exterior points to graph edges, thresholded components, and size-distribution proxies", "Figure 6")

    table_from_rows(
        doc,
        "EdgeConv model and training configuration for the 150-fragment, 100-scene run",
        [
            ["Component", "Configuration"],
            ["Input", "Exterior point coordinates and local graph edges"],
            ["Target", "Balanced positive and negative same-fragment edge labels"],
            ["Backbone", "Dynamic graph EdgeConv-style local feature aggregation"],
            ["Training scenes", "60"],
            ["Validation scenes", "20"],
            ["Held-out test scenes", "20"],
            ["Epochs", "24"],
            ["Validation selection", "Threshold selected on validation scenes before test evaluation"],
            ["Post-processing variants", "Plain thresholding, absorb/merge rules, and post-splitting variants"],
        ],
        [3100, 6260],
        "Table 4",
    )

    heading(doc, "Baseline methods", 2)
    p(doc, (
        "All baseline methods were evaluated on the same 150-fragment scenes and the same 60:20:20 split. The non-learning "
        "baselines used validation-selected hyperparameters. The multilayer perceptron baseline used the same edge-affinity "
        "target but without EdgeConv neighbourhood feature aggregation, which isolates the value of graph-based local "
        "feature learning."
    ))

    heading(doc, "Size-distribution proxy evaluation")
    p(doc, (
        "Predicted connected components were converted to a visible-surface particle-size proxy using a principal-component "
        "span diameter and an equivalent-volume proxy. The eighty per cent passing size was then interpolated from the "
        "volume-weighted cumulative curve. The same conversion was used for every method so that differences in reported "
        "error come from the predicted components rather than from method-specific size calibration."
    ))
    p(doc, (
        "Instance metrics and size metrics were both reported. Normalised mutual information and adjusted Rand index "
        "measure how well the point partition matches fragment labels. Mean absolute passing-size error measures whether "
        "the resulting components recover the coarse size proxy. These two outcomes are related but not identical: a "
        "method can produce imperfect fragment instances while still giving a useful size proxy, or can over-merge and "
        "produce a poor coarse-size estimate despite superficially coherent clusters."
    ))
    heading(doc, "Exterior-to-full PSD calibration", 2)
    p(doc, (
        "Because an exterior scan cannot observe buried fragment volume, the final PSD comparison was separated into three "
        "stages. First, segmentation quality was measured by comparing predicted component labels with exterior fragment "
        "labels. Second, the predicted exterior P80 was compared with an oracle exterior P80 computed from the same retained "
        "HPR points but using the true fragment labels. Third, full-pile PSD recovery was evaluated after a simple linear "
        "calibration from predicted exterior P80 to full ground-truth P80. The calibration model, P80_full = a P80_exterior "
        "+ b, was fitted only on the 20 validation scenes and then frozen before evaluating the 20 test scenes."
    ))
    p(doc, (
        "This calibration layer is not intended to hide segmentation error. Instead, it recognises that even perfect "
        "exterior labels do not reproduce the full-pile size distribution exactly. Reporting the oracle exterior bias, "
        "the predicted exterior-proxy error, and the calibrated full-P80 error makes the visibility error and the model "
        "error explicit."
    ))

    heading(doc, "Results")
    p(doc, (
        "The 24-epoch EdgeConv calibration run showed steadily decreasing training loss and increasing validation average "
        "precision, confirming that the model learned informative same-fragment edge affinities on the HPR exterior target. "
        "Validation average precision rose from 0.905 at epoch 1 to 0.962 at epoch 24. The remaining difficulty was not "
        "edge learning but calibration of the affinity threshold, component merging, and the surface-cluster-to-P80 proxy."
    ))
    add_image(doc, "02_edgeconv_training_curve.png", 14.5, "Training loss and validation average precision for the 24-epoch EdgeConv edge-affinity calibration run on the HPR exterior dataset", "Figure 7")

    edge = read_csv(TABLE_DIR / "edgeconv_test_summary.csv")
    edge_rows = [["Variant", "Thresh.", "Mean absolute P80 error (%)", "Median absolute P80 error (%)", "Mean NMI", "Mean ARI", "Noise"]]
    for r in edge:
        edge_rows.append([
            r["variant"].replace("_", " "),
            fmt(r["threshold"], 4),
            fmt(r["mean_abs_P80_error_pct"]),
            fmt(r["median_abs_P80_error_pct"]),
            fmt(r["mean_NMI"], 3),
            fmt(r["mean_ARI"], 3),
            fmt(r["mean_noise_fraction"], 3),
        ])
    table_from_rows(
        doc,
        "Held-out EdgeConv post-processing comparison for the 20 test scenes",
        edge_rows,
        [2100, 1100, 1500, 1500, 1000, 1000, 1160],
        "Table 5",
    )

    doc.add_page_break()
    bias = read_csv(TABLE_DIR / "hpr_oracle_exterior_to_full_bias_summary.csv")
    bias_rows = [["Split", "Scenes", "Oracle ext. P80 (mm)", "Full P80 (mm)", "Signed bias (%)", "Visible fragments", "Exterior points"]]
    for r in bias:
        bias_rows.append([
            r["split"],
            r["n_scenes"],
            fmt(r["mean_oracle_exterior_P80_mm"]),
            fmt(r["mean_full_ground_truth_P80_mm"]),
            fmt(r["mean_signed_bias_pct"]),
            fmt(r["mean_visible_fragments"], 1),
            fmt(r["mean_exterior_points"], 0),
        ])
    table_from_rows(
        doc,
        "Oracle exterior-label P80 bias relative to the full ground-truth PSD",
        bias_rows,
        [1000, 900, 1700, 1500, 1400, 1400, 1460],
        "Table 6",
    )

    calibrated = read_csv(TABLE_DIR / "edgeconv_three_stage_calibrated_test_summary.csv")
    cal_rows = [["Variant", "NMI", "ARI", "Noise", "Ext. proxy err. vs oracle (%)", "Raw full P80 err. (%)", "Cal. full P80 err. (%)", "Cal. signed err. (%)"]]
    for r in calibrated:
        cal_rows.append([
            r["variant"].replace("edgeconv_", "").replace("_", " "),
            fmt(r["mean_NMI"], 3),
            fmt(r["mean_ARI"], 3),
            fmt(r["mean_noise_fraction"], 3),
            fmt(r["exterior_proxy_mean_abs_error_pct_vs_oracle"]),
            fmt(r["raw_full_mean_abs_P80_error_pct"]),
            fmt(r["calibrated_full_mean_abs_P80_error_pct"]),
            fmt(r["calibrated_full_mean_signed_error_pct"]),
        ])
    table_from_rows(
        doc,
        "Three-stage held-out EdgeConv evaluation: segmentation, exterior-proxy recovery, and validation-calibrated full-P80 recovery",
        cal_rows,
        [1750, 750, 750, 850, 1500, 1300, 1350, 1110],
        "Table 7",
    )

    comp = read_csv(TABLE_DIR / "model_comparison_150frag_test_summary.csv")
    comp_rows = [["Method", "Selected setting", "Mean absolute P80 error (%)", "Median absolute P80 error (%)", "Mean NMI", "Mean ARI", "Noise"]]
    for r in comp:
        method = r["method"]
        if method.startswith("EdgeConv"):
            method = "EdgeConv hybrid bridge"
        comp_rows.append([
            method,
            r["setting"],
            fmt(r["mean_abs_P80_error_pct"]),
            fmt(r["median_abs_P80_error_pct"]),
            fmt(r["mean_NMI"], 3),
            fmt(r["mean_ARI"], 3),
            fmt(r["mean_noise_fraction"], 3),
        ])
    table_from_rows(
        doc,
        "Uncalibrated 150-fragment held-out test comparison after retraining or recalibrating baseline methods on the same scene split",
        comp_rows,
        [1500, 2600, 1300, 1300, 900, 900, 860],
        "Table 8",
    )
    add_image(doc, "model_comparison_150frag_p80.png", 14.0, "Held-out test mean absolute P80 error for the HPR exterior 150-fragment comparison", "Figure 8")
    add_image(doc, "03_edgeconv_test_p80_error_histogram.png", 13.0, "Held-out test distribution of absolute P80 error for the validation-selected EdgeConv hybrid setting", "Figure 9")

    heading(doc, "Discussion")
    p(doc, (
        "The results suggest that the benchmark is diagnosing three separate behaviours. EdgeConv learns useful edge "
        "affinity, as indicated by the training curve and validation average precision. The selected high-instance-quality "
        "hybrid setting also gave substantially stronger instance diagnostics than the MLP baseline (mean NMI 0.754 and "
        "ARI 0.349, compared with NMI 0.134 and negative ARI for MLP). However, raw exterior P80 is not a direct full-pile "
        "PSD estimate. Even the oracle exterior labels overestimated full-pile P80 by 22.94 per cent on the test split, "
        "which explains why uncalibrated component P80 rankings can be misleading."
    ))
    p(doc, (
        "This finding is operationally important. A segmentation model should not be judged only by edge average precision "
        "or only by a final uncalibrated P80 value. If the objective is muckpile monitoring, the post-processing and PSD "
        "stages must be calibrated for the probability distribution produced by the model and for the exterior morphology "
        "of the pile. With the validation-fitted exterior-to-full calibration, the EdgeConv variants achieved 5.00-5.46 "
        "per cent mean absolute full-P80 error on the held-out test split. The calibration does not remove the need for "
        "field validation, but it shows that the raw 29.26 per cent EdgeConv P80 error was largely an exterior-to-full "
        "mapping problem rather than a failure to learn same-fragment geometry."
    ))
    p(doc, (
        "The uncalibrated baseline comparison remains useful as a diagnostic. The MLP affinity baseline produced the lowest "
        "raw P80 error, but it did so with a high noise fraction and weak partition agreement. Density clustering and simple "
        "region growing can capture some surface continuity, but they are brittle when point density, local slopes, and "
        "fragment spacing vary. EdgeConv is therefore retained as the main learning model for fragment-aware exterior "
        "segmentation, while the calibrated result identifies the next bottleneck: linking exterior scan statistics to "
        "field-calibrated full PSD."
    ))

    heading(doc, "Limitations and field validation")
    p(doc, (
        "The benchmark is synthetic and exterior-only. It does not replace field calibration against sieve analysis, belt "
        "scanner measurements, excavated fragment surveys, or other independent references. The pile generator preserves "
        "fragment meshes and labels but is not a full validation of blast breakage, rebound, weathering, fines generation, "
        "or operational digging disturbance. The reported numbers should therefore be interpreted as controlled pre-field "
        "performance rather than mine-site accuracy."
    ))
    p(doc, (
        "Future work should include field point clouds with independent size references, better threshold calibration, "
        "noise-penalised validation objectives, component-merge rules, and contact models that move from axis-clump "
        "approximations toward multi-sphere or polyhedral fragment contact where computation permits."
    ))

    heading(doc, "Conclusions")
    p(doc, (
        "A labelled synthetic exterior point-cloud benchmark was developed for rockpile fragmentation monitoring. The "
        "workflow generates individual fragment meshes, constructs 150-fragment rockpile scenes, filters the full geometry "
        "to exterior-only scan proxies, and evaluates same-fragment edge-affinity learning against multiple baselines. "
        "The latest HPR-refiltered 100-scene comparison showed that EdgeConv learned strong local affinity and produced "
        "better instance diagnostics than the shallow edge-affinity baseline. It also showed that exterior-only PSD "
        "estimation requires calibration: oracle exterior labels overestimated full test P80 by 22.94 per cent, whereas "
        "a validation-fitted exterior-to-full calibration reduced EdgeConv full-P80 error to approximately 5 per cent on "
        "the held-out test scenes. The benchmark therefore provides a useful pre-field environment for separating model "
        "learning, component formation, exterior observation bias, and full-PSD calibration."
    ))

    heading(doc, "Acknowledgements")
    p(doc, (
        "No external funding was received for this study. OpenAI Codex was used as an assistance tool for code refactoring, "
        "document formatting, and drafting support. The author reviewed the scientific content, analysis, and final text."
    ))
    heading(doc, "CRediT author statement")
    p(doc, (
        "N. Ha: Conceptualization, Methodology, Software, Investigation, Validation, Formal analysis, Data curation, "
        "Writing - original draft, Writing - review and editing, Visualization."
    ))
    heading(doc, "Data and code availability")
    p(doc, (
        "The code, synthetic scene metadata, tables, and manuscript assets are available in the project repository: "
        "https://github.com/Linkgyu/rockpile-edgeconv-arxiv. The field validation stage has not yet been completed."
    ))
    heading(doc, "Declaration of competing interest")
    p(doc, "The author declares no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.")

    heading(doc, "References")
    refs = [
        "Cunningham, C.V.B. 1983. The Kuz-Ram model for prediction of fragmentation from blasting. Proceedings of the 1st International Symposium on Rock Fragmentation by Blasting, Lulea, Sweden.",
        "Eltner, A., Kaiser, A., Castillo, C., Rock, G., Neugirg, F., and Abellan, A. 2016. Image-based surface reconstruction in geomorphometry: merits, limits and developments. Earth Surface Dynamics, vol. 4, pp. 359-389.",
        "Engin, I.C., Maerz, N.H., Boyko, K.J., and Reals, R. 2020. Practical measurement of size distribution of blasted rocks using LiDAR scan data. Rock Mechanics and Rock Engineering, vol. 53, pp. 4653-4671.",
        "Ester, M., Kriegel, H.-P., Sander, J., and Xu, X. 1996. A density-based algorithm for discovering clusters in large spatial databases with noise. Proceedings of KDD.",
        "Hubert, L. and Arabie, P. 1985. Comparing partitions. Journal of Classification, vol. 2, pp. 193-218.",
        "Ouchterlony, F. 2005. The Swebrec function: linking fragmentation by blasting and crushing. Mining Technology, vol. 114, pp. 29-44.",
        "Palangio, T.C., Franklin, J.A., and Maerz, N.H. 1995. WIPFRAG: a breakthrough in fragmentation measurement. Proceedings of the 21st Annual Conference on Explosives and Blasting Technique.",
        "Paszke, A., Gross, S., Massa, F., Lerer, A., Bradbury, J., Chanan, G., Killeen, T., Lin, Z., Gimelshein, N., Antiga, L., et al. 2019. PyTorch: an imperative style, high-performance deep learning library. Advances in Neural Information Processing Systems.",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., et al. 2011. Scikit-learn: machine learning in Python. Journal of Machine Learning Research, vol. 12, pp. 2825-2830.",
        "Qi, C.R., Su, H., Mo, K., and Guibas, L.J. 2017a. PointNet: deep learning on point sets for 3D classification and segmentation. Proceedings of CVPR.",
        "Qi, C.R., Yi, L., Su, H., and Guibas, L.J. 2017b. PointNet++: deep hierarchical feature learning on point sets in a metric space. Advances in Neural Information Processing Systems.",
        "Siddiqui, F.I., Shah, S.M.A., and Behan, M.Y. 2009. Measurement of size distribution of blasted rock using digital image processing. Journal of King Abdulaziz University: Engineering Sciences, vol. 20, pp. 81-93.",
        "Vinh, N.X., Epps, J., and Bailey, J. 2010. Information theoretic measures for clusterings comparison: variants, properties, normalisation and correction for chance. Journal of Machine Learning Research, vol. 11, pp. 2837-2854.",
        "Wang, Y., Sun, Y., Liu, Z., Sarma, S.E., Bronstein, M.M., and Solomon, J.M. 2019. Dynamic graph CNN for learning on point clouds. ACM Transactions on Graphics, vol. 38, article 146.",
        "Westoby, M.J., Brasington, J., Glasser, N.F., Hambrey, M.J., and Reynolds, J.M. 2012. Structure-from-Motion photogrammetry: a low-cost, effective tool for geoscience applications. Geomorphology, vol. 179, pp. 300-314.",
    ]
    for ref in refs:
        rp = p(doc, ref)
        rp.paragraph_format.first_line_indent = Cm(-0.5)
        rp.paragraph_format.left_indent = Cm(0.5)

    doc.core_properties.title = title
    doc.core_properties.author = "N. Ha"
    doc.core_properties.comments = "Prepared for Journal of the Southern African Institute of Mining and Metallurgy author guidelines"
    doc.save(OUT_MANUSCRIPT)


def build_cover_letter():
    doc = Document()
    configure_doc(doc)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(3)
    # Remove page footer from the one-page cover letter body if possible.
    for paragraph in doc.sections[0].footer.paragraphs:
        paragraph.text = ""

    p(doc, "23 June 2026")
    p(doc, "Editor")
    p(doc, "Journal of the Southern African Institute of Mining and Metallurgy")
    p(doc)
    p(doc, "Dear Editor,")
    p(doc)
    p(doc, (
        "Please consider the manuscript entitled \"Synthetic exterior point-cloud benchmark for rockpile fragmentation: "
        "fragment generation, rockpile construction, EdgeConv learning, and size-distribution proxy evaluation\" for "
        "publication in the Journal of the Southern African Institute of Mining and Metallurgy."
    ))
    p(doc, (
        "The manuscript is aligned with the Journal's mining, blasting, digitalisation, and mine-to-mill interests. It "
        "presents a reproducible synthetic benchmark in which rock-like fragments are generated first, assembled into "
        "labelled rockpile scenes, converted to exterior-only point-cloud scan proxies using multi-view Hidden Point "
        "Removal, and used to evaluate EdgeConv edge-affinity learning and several baseline clustering methods for visible-surface size-distribution proxy "
        "estimation."
    ))
    p(doc, (
        "The main novelty is the separation of labelled pile construction, exterior-only scan visibility, and final "
        "particle-size proxy recovery. The latest 100-scene, 150-fragment benchmark uses a 60:20:20 train, validation, "
        "and test split and compares EdgeConv with graph-threshold, density, region-growing, and multilayer perceptron "
        "edge-affinity baselines. The revised HPR experiment separates segmentation performance, exterior-proxy P80 "
        "performance, and validation-calibrated full-PSD performance. This separation shows that EdgeConv learns stronger "
        "fragment-aware affinity than the shallow baseline and that exterior-to-full PSD calibration is required because "
        "even oracle exterior labels are biased relative to the full pile."
    ))
    p(doc, (
        "The work is original, has not been published previously, and is not under consideration by another journal. The "
        "author approves the submission and declares no competing financial interests or personal relationships that could "
        "have appeared to influence the work. No external funding was received. The manuscript discloses the use of OpenAI "
        "Codex for code refactoring, document formatting, and drafting assistance; all scientific decisions and final text "
        "were reviewed by the author."
    ))
    p(doc, (
        "Correspondence should be addressed to N. Ha, Department of Energy and Resources Engineering, National Korea "
        "Maritime and Ocean University, Busan, Republic of Korea; Email: noah.ha.au@gmail.com; ORCID: 0009-0007-4918-6239."
    ))
    p(doc, "Yours sincerely,")
    p(doc, "N. Ha")

    doc.core_properties.title = "Cover letter - SAIMM submission"
    doc.core_properties.author = "N. Ha"
    doc.save(OUT_COVER)


if __name__ == "__main__":
    build_manuscript()
    build_cover_letter()
    print(OUT_MANUSCRIPT)
    print(OUT_COVER)
